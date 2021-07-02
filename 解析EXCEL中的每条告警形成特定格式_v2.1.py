#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Author : N0Coriander
# address : https://github.com/N0Coriander
# Date : 2021/6/26 13:44
# Desc : 用户自己先针对EXCEL进行筛选，选中要生成通告的数据(支持多行)，然后根据系统提示生成以下五类通告
# Compare : 上报通告自动追加进word文档，实现了HW日报的自动化输出
# --------------------------------------------------------
# 成功类———按事件名称归类，通告格式用最全的那个
# 爆破类———按事件名称归类，罗列汇总的攻击、受害IP，增加次数
# 远程工具———按事件名称归类，只统计受害IP
# 外打内———事件名称统一输出为"存在扫描行为"，只统计攻击者IP
# 内打内———事件名称统一输出为"内网IP攻击行为"，罗列攻击者IP、受害IP
# --------------------------------------------------------

# 自动下载并导入pyperclip和docx
try:
    import os
    import pyperclip
    from docx import Document
except ImportError:
    import os

    os.system('pip3 install pyperclip')
    os.system('pip3 install python-docx')
    import pyperclip
    from docx import Document

seq = 0
while 1:
    input('[+] 请先在EXCEL中筛选你要生成通告的告警，然后选中这些行并复制，准备好后按回车键开始')

    info = pyperclip.paste()  # 读取剪贴板内容
    # 剪贴板不为空判断
    if info:
        split_info = info.split('\r\n')
    else:
        print('[ERROR] 大哥,你还没复制呢！')
        continue

    # 用于存放EXCEL中的每一行
    excel_line = []
    for i in split_info:
        end = i.split('\t')
        excel_line.append(end)
    for q in excel_line:
        if q == ['']:
            excel_line.remove([''])
        else:
            pass

    # 用于存放受害IP
    victim_ip = []
    for h in excel_line:
        victim_ip.append(h[2])
    victim_ip_rtd = list(set(victim_ip))  # 受害IP去重
    for q in victim_ip_rtd:
        if q == ['']:
            victim_ip_rtd.remove([''])
        else:
            pass

    # 用于存放攻击者IP
    attack_ip = []
    for v in excel_line:
        attack_ip.append(v[3])
    attack_ip_rtd = list(set(attack_ip))  # 攻击者IP去重
    for q in attack_ip_rtd:
        if q == ['']:
            attack_ip_rtd.remove([''])
        else:
            pass

    # 用于存放事件名称
    event_name = []
    for k in excel_line:
        event_name.append(k[5])
    event_name_rtd = list(set(event_name))  # 事件名称去重
    for q in event_name_rtd:
        if q == ['']:
            event_name_rtd.remove([''])
        else:
            pass

    while 1:
        # 获取用户输入
        user_input = input('[+] 根据类别输入相应数字（ 成功类:1 / 爆破:2 / 远程工具:3 / 外打内:4 / 内打内:5 ）: ')
        # 列出目录下的所有文件，把word文件名放到一起
        files = os.listdir()
        word_file = []
        for i in files:
            if 'docx' in i:
                word_file.append(i)
            else:
                pass

        # 定义document为空，是因为发现直接走下面的循环之后，系统报错说"NameError: name 'document' is not defined"
        document = ''

        # 适配新建或追加两种场景，所以使用了for循环+if判断
        if len(word_file) > 0:
            for i in word_file:
                if (excel_line[0][1] + '-通告.docx') == i:
                    document = Document(f'{excel_line[0][1]}-通告.docx')
                else:
                    document = Document()
        else:
            document = Document()

        # 成功类———按事件名称分，一样名字的聚合到一起，通告格式用最全的那个
        if user_input == '1':
            end = []
            # 遍历去重后的告警名称
            for j in event_name_rtd:
                seq = seq + 1
                attack_ip_1 = []
                victim_ip_1 = []
                time = []
                # 遍历去重前的每一行EXCEL数据
                for g in excel_line:
                    if j == g[5]:
                        time.append(g[0])  # 将告警时间追加到列表中，最后我只要第一个告警时间，视为最新
                        attack_ip_1.append(g[3])  # 提取告警名称一样的攻击IP
                        victim_ip_1.append(g[2])  # 提取告警名称一样的受害IP
                attack_ip_1_rtd = list(set(attack_ip_1))
                victim_ip_1_rtd = list(set(victim_ip_1))

                tonggao1 = '事件名称：' + j + '\n'
                tonggao2 = '客户：' + excel_line[0][1] + '\n'
                tonggao3 = '报告序号：' + str(seq) + '\n'
                tonggao4 = '最近告警时间：' + time[0] + '\n'
                tonggao5 = '攻击者IP：' + '、'.join(attack_ip_1_rtd) + '\n'
                tonggao6 = '受害IP：' + '、'.join(victim_ip_1_rtd) + '\n'
                tonggao7 = '研判结论：攻击成功/登录成功/脆弱性/...\n'
                tonggao8 = '建议：在防火墙等边界设备上封禁攻击IP/确认是否为业务触发，若不是立即下线并上机排查/强化口令xxx' + '\n'
                tonggao = tonggao1 + tonggao2 + tonggao3 + tonggao4 + tonggao5 + tonggao6 + tonggao7 + tonggao8
                end.append(tonggao)  # 以事件名称分类，将该通告追加到空列表中，方便最后一起粘贴
                document.add_heading(tonggao1, 1)
                document.add_paragraph(tonggao2 + tonggao3 + tonggao4 + tonggao5 + tonggao6 + tonggao7 + tonggao8)
                document.save(f'{excel_line[0][1]}-通告.docx')

        # 爆破类———按威胁名称归类，罗列攻击者IP、受害IP、次数
        elif user_input == '2':
            end = []
            # 遍历去重后的告警名称
            for j in event_name_rtd:
                seq = seq + 1
                attack_ip_1 = []
                victim_ip_1 = []
                nums = []   # 次数
                # 遍历去重前的每一行EXCEL数据
                for g in excel_line:
                    if j == g[5]:
                        nums.append(g[9])
                        attack_ip_1.append(g[3])  # 提取告警名称一样的攻击IP
                        victim_ip_1.append(g[2])  # 提取告警名称一样的受害IP
                num = 0
                for nb in nums:
                    num = num + int(nb)
                attack_ip_1_rtd = list(set(attack_ip_1))
                victim_ip_1_rtd = list(set(victim_ip_1))

                tonggao1 = '事件名称：' + j + '\n'
                tonggao2 = '客户：' + excel_line[0][1] + '\n'
                tonggao3 = '报告序号：' + str(seq) + '\n'
                tonggao5 = '攻击者IP：' + '、'.join(attack_ip_1_rtd) + '\n'
                tonggao6 = '受害IP：' + '、'.join(victim_ip_1_rtd) + '\n'
                tonggao7 = '次数：' + str(num) + '\n'
                tonggao8 = '建议：确认是否为业务频繁登录触发，若不是立即下线并上机排查' + '\n'
                tonggao = tonggao1 + tonggao2 + tonggao3 + tonggao5 + tonggao6 + tonggao7 + tonggao8
                end.append(tonggao)  # 以事件名称分类，将该通告追加到空列表中，方便最后一起粘贴
                document.add_heading(tonggao1, 1)
                document.add_paragraph(tonggao2 + tonggao3 + tonggao5 + tonggao6 + tonggao7 + tonggao8)
                document.save(f'{excel_line[0][1]}-通告.docx')

        # 远程工具———只统计受害IP
        elif user_input == '3':
            end = []
            # 遍历去重后的告警名称
            for j in event_name_rtd:
                seq = seq + 1
                victim_ip_1 = []
                # 遍历去重前的每一行EXCEL数据
                for g in excel_line:
                    if j == g[5]:
                        victim_ip_1.append(g[2])  # 提取告警名称一样的受害IP
                victim_ip_1_rtd = list(set(victim_ip_1))

                tonggao1 = '事件名称：' + j + '\n'
                tonggao2 = '客户：' + excel_line[0][1] + '\n'
                tonggao3 = '报告序号：' + str(seq) + '\n'
                tonggao6 = '受害IP：' + '、'.join(victim_ip_1_rtd) + '\n'
                tonggao8 = '建议：确认是否为正常业务使用' + '\n'
                tonggao = tonggao1 + tonggao2 + tonggao3 + tonggao6 + tonggao8
                end.append(tonggao)  # 以事件名称分类，将该通告追加到空列表中，方便最后一起粘贴
                document.add_heading(tonggao1, 1)
                document.add_paragraph(tonggao2 + tonggao3 + tonggao6 + tonggao8)
                document.save(f'{excel_line[0][1]}-通告.docx')

        # 外打内———事件名称统一输出为"存在扫描行为"，只统计攻击者IP
        elif user_input == '4':
            end = []
            seq = seq + 1
            tonggao1 = '事件名称：存在扫描行为' + '\n'
            tonggao2 = '客户：' + excel_line[0][1] + '\n'
            tonggao3 = '报告序号：' + str(seq) + '\n'
            tonggao5 = '攻击者IP：' + '、'.join(attack_ip_rtd) + '\n'
            tonggao8 = '建议：在防火墙等边界设备上封禁' + '\n'
            tonggao = tonggao1 + tonggao2 + tonggao3 + tonggao5 + tonggao8
            end.append(tonggao)  # 以事件名称分类，将该通告追加到空列表中，方便最后一起粘贴
            document.add_heading(tonggao1, 1)
            document.add_paragraph(tonggao2 + tonggao3 + tonggao5 + tonggao8)
            document.save(f'{excel_line[0][1]}-通告.docx')

        # 内打内———事件名称统一输出为"内网IP攻击行为"，罗列攻击者IP、受害IP
        elif user_input == '5':
            end = []
            seq = seq + 1
            tonggao1 = '事件名称：内网IP攻击行为' + '\n'
            tonggao2 = '客户：' + excel_line[0][1] + '\n'
            tonggao3 = '报告序号：' + str(seq) + '\n'
            tonggao5 = '攻击者IP：' + '、'.join(attack_ip_rtd) + '\n'
            tonggao6 = '受害IP：' + '、'.join(victim_ip_rtd) + '\n'
            tonggao8 = '建议：确认是否为流量转发或漏扫类设备，若不是，建议立即上机排查' + '\n'
            tonggao = tonggao1 + tonggao2 + tonggao3 + tonggao5 + tonggao6 + tonggao8
            end.append(tonggao)  # 以事件名称分类，将该通告追加到空列表中，方便最后一起粘贴
            document.add_heading(tonggao1, 1)
            document.add_paragraph(tonggao2 + tonggao3 + tonggao5 + tonggao6 + tonggao8)
            document.save(f'{excel_line[0][1]}-通告.docx')

        else:
            print('[WARNING] 瞎几把输！')
            continue

        for q in end:
            if q == ['']:
                end.remove([''])
            else:
                pass
        pyperclip.copy('\n'.join(end))  # 换行分割每一个通告

        break
    print('[!] 内容已提取到剪贴板，请自主粘贴')
